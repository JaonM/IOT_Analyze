# -*- coding:utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from analyze import analyze
import datetime
import os
import re
from visualization.keep import draw_picture
from analyze.hotel_sensor_config import sensor_config

input_path = '../input/'


# sensor_config = {
#     '中科能源': [
#         {'eui': '3786E6ED0034004B', 'frequency': 300},
#         {'eui': '3768B26900230053', 'frequency': 300},
#         {'eui': '4768B269002B0059', 'frequency': 300},
#         {'eui': '4768B269001F003F', 'frequency': 300},
#         {'eui': '4778B269003B002F', 'frequency': 300}
#     ],
#     '安盈科技': [
#         {'eui': '9896830000000008', 'frequency': 300},
#         {'eui': '9896830000000002', 'frequency': 300},
#         {'eui': '9896830000000003', 'frequency': 300},
#         {'eui': '9896830000000004', 'frequency': 300},
#         {'eui': '9896830000000006', 'frequency': 300}
#     ],
#     '伟曼达': [
#         {'eui': '3430363057376506', 'frequency': 900},
#         {'eui': '3430363064377B07', 'frequency': 300},
#         {'eui': '3430363064378607', 'frequency': 300},
#         {'eui': '343036305D375E05', 'frequency': 300},
#         {'eui': '3430363064378007', 'frequency': 900}
#     ]
# }


def create_report(date=datetime.datetime.now().date().strftime('%Y-%m-%d'), total_result=None):
    if total_result is None:
        print('analyze result is not ready yet')
        return
    print('generating the report please wait...')

    dimensions = ['温度', '湿度', '电流', '电压', '功率', '电量', '时间频率']
    # main page
    for key in sensor_config.keys():
        document = Document()

        # 基地报告生成
        # add a cover page title
        document.add_paragraph()
        title = document.add_paragraph(key + ' 丽思卡尔顿传感器分析报告' + ' ' + date)
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_before = Pt(250)
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.name = 'Simhei'

        document.add_page_break()
        for result in total_result['results']:
            if result['company'] == key:
                # 基地传感器
                # if result['is_base'] == 1:
                document.add_heading(result['eui'], level=1)
                document.add_paragraph().paragraph_format.space_before = Pt(4)
                try:
                    document.add_picture(input_path + result['eui'] + '.png')
                except FileNotFoundError:
                    pass
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                if result['info'] == '缺失当天数据':
                    p = document.add_paragraph("缺失当天数据")
                    p.paragraph_format.left_indent = Inches(0.5)
                    continue
                pics = os.listdir('../' + date)
                target_pic = None
                for pic in pics:
                    if re.search(result['eui'], pic):
                        target_pic = pic
                        break
                if target_pic is not None:
                    target_pic = '../' + date + '/' + target_pic
                    # p.alignment = WD_ALIGN_PARAGRAPH
                    p.add_run().add_picture(target_pic, width=Inches(6.5))

                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.aligment = WD_ALIGN_PARAGRAPH

                for dim in dimensions:
                    try:
                        for msg in result['messages'][dim]:
                            if re.search('nan', msg['msg']):
                                continue
                            p.paragraph_format.left_indent = Inches(0.5)

                            run = p.add_run(msg['msg'])
                            if msg['status'] == 'red':  # DC143C
                                run.font.color.rgb = RGBColor(0xDC, 0x14, 0x3C)
                            elif msg['status'] == 'blue':
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
                    except KeyError:
                        continue
                document.add_page_break()
        document.save(date + ' ' + key + '丽思卡尔顿传感器分析报告' + '.docx')


if __name__ == '__main__':
    date_range = ['2018-03-08']
    for date in date_range:
        print("analyzing data please wait ...")
        result = analyze.analyze_data(date=date)
        print('analyzing finished ' + str(result))

        print("drawing pictures..")
        # draw_picture(date=date)
        print("drawing pictures finished")
        create_report(date, result)
    print('process finished')
